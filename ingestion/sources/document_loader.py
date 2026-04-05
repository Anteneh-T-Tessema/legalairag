from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pdfplumber

from config.logging import get_logger

logger = get_logger(__name__)


@dataclass
class ParsedDocument:
    source_id: str
    filename: str
    mime_type: str
    raw_text: str
    pages: list[str]
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return self.raw_text


def load_from_bytes(
    content: bytes,
    source_id: str,
    filename: str,
    metadata: dict[str, Any] | None = None,
) -> ParsedDocument:
    """
    Parse a document from raw bytes.
    Dispatches to the appropriate parser based on file extension / magic bytes.
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf" or content[:4] == b"%PDF":
        return _parse_pdf(content, source_id, filename, metadata or {})
    if suffix in {".doc", ".docx"}:
        return _parse_docx(content, source_id, filename, metadata or {})
    if suffix in {".txt", ".text"}:
        return _parse_plaintext(content, source_id, filename, metadata or {})
    if suffix in {".htm", ".html"}:
        return _parse_html(content, source_id, filename, metadata or {})

    logger.warning("unsupported_format", filename=filename, suffix=suffix)
    text = content.decode("utf-8", errors="replace")
    return ParsedDocument(
        source_id=source_id,
        filename=filename,
        mime_type="text/plain",
        raw_text=text,
        pages=[text],
        metadata=metadata or {},
    )


def _parse_pdf(
    content: bytes,
    source_id: str,
    filename: str,
    metadata: dict[str, Any],
) -> ParsedDocument:
    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages.append(_normalize_whitespace(text))

    full_text = "\n\n".join(pages)
    logger.info("parsed_pdf", source_id=source_id, pages=len(pages), chars=len(full_text))
    return ParsedDocument(
        source_id=source_id,
        filename=filename,
        mime_type="application/pdf",
        raw_text=full_text,
        pages=pages,
        metadata=metadata,
    )


def _parse_docx(
    content: bytes,
    source_id: str,
    filename: str,
    metadata: dict[str, Any],
) -> ParsedDocument:
    from docx import Document  # type: ignore[import-untyped]

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    return ParsedDocument(
        source_id=source_id,
        filename=filename,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        raw_text=full_text,
        pages=[full_text],
        metadata=metadata,
    )


def _parse_plaintext(
    content: bytes,
    source_id: str,
    filename: str,
    metadata: dict[str, Any],
) -> ParsedDocument:
    text = content.decode("utf-8", errors="replace")
    return ParsedDocument(
        source_id=source_id,
        filename=filename,
        mime_type="text/plain",
        raw_text=text,
        pages=[text],
        metadata=metadata,
    )


def _parse_html(
    content: bytes,
    source_id: str,
    filename: str,
    metadata: dict[str, Any],
) -> ParsedDocument:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = _normalize_whitespace(soup.get_text(separator="\n"))
    return ParsedDocument(
        source_id=source_id,
        filename=filename,
        mime_type="text/html",
        raw_text=text,
        pages=[text],
        metadata=metadata,
    )


def _normalize_whitespace(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()
